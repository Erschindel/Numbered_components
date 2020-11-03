from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


components_doc = Document("docs/Components.docx")
specs_doc =  Document("docs/Specs.docx")


# create component and spec lists

def doc_to_list(doc):
    list_ = [i.text.replace("\t", "").lower() for i in doc.paragraphs]

    for i in list_:
        try:
            if i == "" or isinstance(int(i.split(" ")[-1]), int):
                list_.remove(i)
        except:
            pass

    return list_

components_list = doc_to_list(components_doc)
specs_list = [i.text for i in specs_doc.paragraphs]

full_doc = '\t%s' % "\n\t".join(specs_list)


# add numbered components to full_doc

for i, component in enumerate(components_list):
    full_doc = full_doc.replace(component, f"{component} {str(i + 1)}")


# split document into desired runs

split_doc = []
run = []

for letter in full_doc:
    try:
        if isinstance(int(letter), int):
            split_doc.append(("").join(run))
            split_doc.append(letter)
            run = []
    except:
        run.append(letter)

split_doc.append(("").join(run))


# make final document and add styles

document = Document()
doc_style = document.styles["Normal"]
doc_style.font.size = Pt(12)
doc_style.font.name = "Times New Roman"

paragraph = document.add_paragraph(split_doc[0])
paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

for runs in split_doc[1:]:
    try:
        if isinstance(int(runs), int):
            run = paragraph.add_run(runs)
            run.bold = True
    except:
        run = paragraph.add_run(runs)
        run.bold = False

document.save("Specs_numbered.docx")
